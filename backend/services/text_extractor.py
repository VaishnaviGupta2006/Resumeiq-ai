import os
from docx import Document
import pdfplumber

class UnsupportedFileTypeError(Exception):
    """Raised when file type is not supported."""
    pass

class TextExtractionError(Exception):
    """Raised when text extraction fails."""
    pass

def extract_text(file_path):
    """
    Extract text from a resume file (PDF or DOCX).
    
    Args:
        file_path (str): Path to the resume file
        
    Returns:
        str: Extracted text content
        
    Raises:
        UnsupportedFileTypeError: If file type is not supported
        TextExtractionError: If text extraction fails
    """
    if not os.path.exists(file_path):
        raise TextExtractionError(f"File not found: {file_path}")
    
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return _extract_from_pdf(file_path)
    elif file_ext in ['.docx', '.doc']:
        return _extract_from_docx(file_path)
    else:
        raise UnsupportedFileTypeError(f"Unsupported file type: {file_ext}")

def _extract_from_pdf(file_path):
    """Extract text from PDF using pdfplumber."""
    try:
        text = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        
        if not text:
            raise TextExtractionError("No text could be extracted from PDF")
        
        return '\n'.join(text)
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")

def _extract_from_docx(file_path):
    """Extract text from DOCX using python-docx."""
    try:
        doc = Document(file_path)
        text = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)
        
        if not text:
            raise TextExtractionError("No text could be extracted from DOCX")
        
        return '\n'.join(text)
    except Exception as e:
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")
